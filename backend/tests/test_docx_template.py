import pytest
import tempfile
import os
import sys
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate

# Add parent directory to path for imports when run directly
if __name__ == "__main__":
    sys.path.append(str(Path(__file__).parent.parent.parent))

from backend.docx_utils import tag_based_template_replacement, extract_template_tags


class TestDocxTemplate:
    """Test docx template tag replacement functionality."""
    
    def create_test_template(self, template_path: Path, content_with_tags: str) -> None:
        """Create a test docx template with template tags."""
        doc = Document()
        doc.add_paragraph(content_with_tags)
        doc.save(template_path)
    
    def create_test_template_with_docxtpl(self, template_path: Path) -> None:
        """Create a proper docxtpl template for testing."""
        doc = Document()
        
        # Add header content
        doc.add_paragraph("Jerry Mindek")
        doc.add_paragraph("614-560-5114 | jerry.mindek@gmail.com | Columbus, OH")
        doc.add_paragraph("")
        
        # Add template tags
        doc.add_paragraph("{{company}}")
        doc.add_paragraph("Re: {{role}}")
        doc.add_paragraph("{{content}}")
        
        doc.save(template_path)
    
    def test_extract_template_tags_basic(self):
        """Test extracting template tags from Claude's response."""
        mock_llm_response = """{{company}}
TechCorp Inc

{{role}}  
Senior Software Engineer

{{content}}
Dear Hiring Manager,

I am excited to apply for the Senior Software Engineer position at TechCorp Inc. With over 8 years of experience leading engineering teams and building scalable systems, I am confident I would be a valuable addition to your team.

At my previous role at Root Inc, I led the technical architecture for high-scale insurance processing systems that handled millions of transactions. I also managed a team of 6 engineers while maintaining hands-on technical contributions.

I am particularly drawn to TechCorp's mission of innovation and would welcome the opportunity to discuss how my experience can contribute to your continued success.

Sincerely,

Jerry Mindek"""
        
        # Test template tag extraction
        extracted_tags = extract_template_tags(mock_llm_response)
        
        # Verify all tags were extracted
        assert 'company' in extracted_tags
        assert 'role' in extracted_tags  
        assert 'content' in extracted_tags
        
        # Verify content is correct
        assert extracted_tags['company'] == 'TechCorp Inc'
        assert extracted_tags['role'] == 'Senior Software Engineer'
        assert 'Dear Hiring Manager,' in extracted_tags['content']
        assert 'Sincerely,\n\nJerry Mindek' in extracted_tags['content']
        
        print(f"✅ Extracted {len(extracted_tags)} template tags successfully")
    
    def test_extract_template_tags_with_brackets(self):
        """Test extracting template tags that have content wrapped in brackets."""
        mock_llm_response_with_brackets = """{{company}}
[Acme Corporation]

{{role}}
[Engineering Manager]

{{content}}
[Dear Hiring Manager,

I am writing to express my interest in the Engineering Manager position at Acme Corporation.

Sincerely,

Jerry Mindek]"""
        
        extracted_tags = extract_template_tags(mock_llm_response_with_brackets)
        
        # Verify brackets are removed from content
        assert extracted_tags['company'] == 'Acme Corporation'
        assert extracted_tags['role'] == 'Engineering Manager'
        assert extracted_tags['content'] == 'Dear Hiring Manager,\n\nI am writing to express my interest in the Engineering Manager position at Acme Corporation.\n\nSincerely,\n\nJerry Mindek'
        
        print("✅ Successfully handled bracketed content")
    
    def test_docx_template_replacement_integration(self):
        """Test complete integration of template tag replacement in docx files."""
        # Create temporary files
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "test_template.docx"
            output_path = Path(temp_dir) / "test_output.docx"
            
            # Create test template
            self.create_test_template_with_docxtpl(template_path)
            
            # Mock LLM response
            mock_llm_response = """{{company}}
DataCorp Solutions

{{role}}
Lead Data Engineer  

{{content}}
Dear Hiring Manager,

I am excited to apply for the Lead Data Engineer position at DataCorp Solutions. My experience at Root Inc building scalable data processing systems that handled 10M+ daily transactions makes me well-suited for this role.

I led the migration of our legacy data pipeline to a modern cloud-based architecture, reducing processing time by 60% and improving system reliability to 99.9% uptime.

I would welcome the opportunity to discuss how my expertise in data engineering and team leadership can contribute to DataCorp's continued growth.

Sincerely,

Jerry Mindek"""
            
            # Extract tags from mock response
            tag_values = extract_template_tags(mock_llm_response)
            
            # Test the template replacement
            success = tag_based_template_replacement(template_path, output_path, tag_values)
            
            # Verify the operation succeeded
            assert success, "Template replacement should succeed"
            assert output_path.exists(), "Output file should be created"
            
            # Read the output file and verify content was replaced
            output_doc = Document(output_path)
            output_text = '\n'.join([paragraph.text for paragraph in output_doc.paragraphs])
            
            # Verify template tags were replaced with actual content
            assert 'DataCorp Solutions' in output_text
            assert 'Lead Data Engineer' in output_text
            assert 'Dear Hiring Manager,' in output_text
            assert 'Jerry Mindek' in output_text
            
            # Verify template tags are no longer present
            assert '{{company}}' not in output_text
            assert '{{role}}' not in output_text
            assert '{{content}}' not in output_text
            
            print("✅ Complete docx template replacement integration test passed")
    
    def test_tag_replacement_with_missing_tags(self):
        """Test behavior when some template tags are missing from LLM response."""
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "test_template.docx"
            output_path = Path(temp_dir) / "test_output.docx"
            
            # Create test template
            self.create_test_template_with_docxtpl(template_path)
            
            # Mock LLM response missing the 'role' tag
            incomplete_response = """{{company}}
TestCorp Inc

{{content}}
Dear Hiring Manager,

Test content here.

Sincerely,

Jerry Mindek"""
            
            tag_values = extract_template_tags(incomplete_response)
            
            # Should have company and content, but not role
            assert 'company' in tag_values
            assert 'content' in tag_values
            assert 'role' not in tag_values
            
            # Template replacement should still work (docxtpl handles missing tags)
            success = tag_based_template_replacement(template_path, output_path, tag_values)
            assert success, "Should succeed even with missing tags"
            
            print("✅ Handled missing template tags gracefully")


if __name__ == "__main__":
    # Run tests manually if called directly
    test = TestDocxTemplate()
    test.test_extract_template_tags_basic()
    test.test_extract_template_tags_with_brackets()  
    test.test_docx_template_replacement_integration()
    test.test_tag_replacement_with_missing_tags()
    print("🎉 All tests passed!")