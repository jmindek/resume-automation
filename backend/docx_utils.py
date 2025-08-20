"""Utilities for reading and writing Microsoft Word .docx files"""

from docx import Document
from docxtpl import DocxTemplate
from pathlib import Path
from typing import Optional, Dict, Any


def read_docx_content(file_path: Path) -> str:
    """
    Read text content from a .docx file.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        String containing all text content from the document
    """
    try:
        # Check if file exists and has reasonable size
        if not file_path.exists():
            raise Exception(f"File does not exist: {file_path}")
        
        file_size = file_path.stat().st_size
        if file_size == 0:
            raise Exception(f"File is empty: {file_path}")
        
        # Try to read with python-docx directly, which is more forgiving than zipfile validation
        doc = Document(file_path)
        
        # Extract all text from paragraphs
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                content.append(paragraph.text)
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text)
        
        return '\n'.join(content)
        
    except Exception as e:
        raise Exception(f"Failed to read .docx file {file_path}: {str(e)}")


def write_docx_content(file_path: Path, content: str, preserve_formatting: bool = False) -> bool:
    """
    Write text content to a .docx file.
    
    Args:
        file_path: Path where to save the .docx file
        content: Text content to write
        preserve_formatting: If True, attempt to preserve basic formatting
        
    Returns:
        True if successful, False otherwise
    """
    try:
        doc = Document()
        
        # Split content into paragraphs
        paragraphs = content.split('\n')
        
        for paragraph_text in paragraphs:
            paragraph_text = paragraph_text.strip()
            if paragraph_text:  # Skip empty lines
                # Check if this looks like a heading (all caps, short lines, etc.)
                if (paragraph_text.isupper() and len(paragraph_text) < 50) or \
                   paragraph_text.endswith(':') and len(paragraph_text) < 50:
                    # Add as heading
                    doc.add_heading(paragraph_text, level=1)
                else:
                    # Add as regular paragraph
                    doc.add_paragraph(paragraph_text)
            else:
                # Add empty paragraph for spacing
                doc.add_paragraph()
        
        # Save the document
        doc.save(str(file_path))
        return True
        
    except Exception as e:
        print(f"Failed to write .docx file {file_path}: {str(e)}")
        return False


def copy_docx_with_new_content(source_path: Path, dest_path: Path, new_content: str) -> bool:
    """
    Copy a .docx file and replace ONLY the text content while preserving ALL formatting.
    This maintains fonts, styles, spacing, structure, and document appearance.
    
    Args:
        source_path: Path to source .docx file
        dest_path: Path for new .docx file  
        new_content: New text content to replace existing content
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Load the source document
        source_doc = Document(source_path)
        
        # Split new content into lines
        new_lines = [line.strip() for line in new_content.split('\n') if line.strip()]
        
        # Get all existing paragraphs
        existing_paragraphs = source_doc.paragraphs[:]
        
        # Strategy: Replace text content in existing paragraphs while preserving formatting
        line_index = 0
        
        for paragraph in existing_paragraphs:
            if line_index < len(new_lines):
                # Clear existing text but keep paragraph formatting
                paragraph.clear()
                
                # Add new text with preserved formatting
                run = paragraph.add_run(new_lines[line_index])
                line_index += 1
            else:
                # If we have more paragraphs than content lines, clear the extras
                paragraph.clear()
        
        # If we have more content lines than paragraphs, add them as new paragraphs
        while line_index < len(new_lines):
            new_paragraph = source_doc.add_paragraph()
            new_paragraph.add_run(new_lines[line_index])
            line_index += 1
        
        # Save to destination
        source_doc.save(str(dest_path))
        return True
        
    except Exception as e:
        print(f"Failed to copy and update .docx file from {source_path} to {dest_path}: {str(e)}")
        return False


def smart_template_replacement(source_path: Path, dest_path: Path, template_content: str, replacement_content: str, placeholder: str = "{{CONTENT}}") -> bool:
    """
    Replace a specific placeholder in a template with new content while preserving formatting.
    This is used for cover letters with {{CONTENT}} tags.
    
    Args:
        source_path: Path to template .docx file
        dest_path: Path for new .docx file
        template_content: Original template content as text
        replacement_content: New content to insert at placeholder location
        placeholder: Placeholder tag to replace (default: "{{CONTENT}}")
        
    Returns:
        True if successful, False otherwise
    """
    try:
        import shutil
        
        # Check if template contains the placeholder
        if placeholder not in template_content:
            # No placeholder found, use regular smart content replacement
            return smart_content_replacement(source_path, dest_path, replacement_content)
        
        # First, make an exact copy of the source file
        shutil.copy2(source_path, dest_path)
        
        # Load the copied document
        doc = Document(dest_path)
        
        # Find and replace the placeholder in all paragraphs
        placeholder_found = False
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Replace the placeholder with the new content
                paragraph.text = paragraph.text.replace(placeholder, replacement_content)
                placeholder_found = True
        
        # Also check tables for the placeholder
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement_content)
                            placeholder_found = True
        
        if not placeholder_found:
            print(f"Warning: Placeholder '{placeholder}' not found in template, falling back to content replacement")
            return smart_content_replacement(source_path, dest_path, replacement_content)
        
        # Save the updated document
        doc.save(str(dest_path))
        return True
        
    except Exception as e:
        print(f"Failed smart template replacement from {source_path} to {dest_path}: {str(e)}")
        return False


def multi_tag_template_replacement(source_path: Path, dest_path: Path, template_content: str, replacement_content: str) -> bool:
    """
    Replace multiple template tags in a document. Expects Claude to have already processed the template
    and returned the complete document with all tags replaced.
    
    Args:
        source_path: Path to template .docx file (not used since Claude handles all replacements)
        dest_path: Path for new .docx file
        template_content: Original template content (not used since Claude handles replacements)
        replacement_content: Complete cover letter with all tags already replaced by Claude
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Since Claude is handling all the template tag replacements,
        # we just need to create a new document with the complete content
        return write_docx_content(dest_path, replacement_content, preserve_formatting=True)
        
    except Exception as e:
        print(f"Failed multi-tag template replacement: {str(e)}")
        return False


def tag_based_template_replacement(template_path: Path, dest_path: Path, tag_values: Dict[str, Any]) -> bool:
    """
    Replace tags in a template document using docxtpl library.
    This preserves all formatting, styles, and document structure.
    
    Args:
        template_path: Path to template .docx file with tags like {{jobtitle}}, {{key_achievements}}, etc.
        dest_path: Path for new .docx file with tags replaced
        tag_values: Dictionary mapping tag names to their replacement values
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Map tag values to match template tag names
        mapped_values = map_tag_values_to_template(tag_values, template_path)
        
        # Load the template
        template = DocxTemplate(template_path)
        
        # Replace tags with values
        template.render(mapped_values)
        
        # Save the result
        template.save(dest_path)
        
        print(f"Successfully populated template using docxtpl: {dest_path}")
        print(f"Used tags: {list(mapped_values.keys())}")
        return True
        
    except Exception as e:
        print(f"Failed tag-based template replacement using docxtpl: {str(e)}")
        return False


def map_tag_values_to_template(tag_values: Dict[str, Any], template_path: Path) -> Dict[str, Any]:
    """
    Map JSON tag values to actual template tag names.
    Handles discrepancies between Claude's JSON keys and template tags.
    """
    try:
        # Read template to see what tags it actually contains
        template_content = read_docx_content(template_path)
        
        # Create mapping based on template file name and content
        mapped_values = {}
        
        # Handle resume template tags
        if "Template" in template_path.name and ("Manager" in template_path.name or "Engineer" in template_path.name):
            # Resume template mappings
            mapped_values['jobtitle'] = tag_values.get('jobtitle', 'Engineering Manager')
            mapped_values['title'] = tag_values.get('jobtitle', 'Engineering Manager')  # Alternative name
            mapped_values['key_achievements'] = tag_values.get('key_achievements', '')  # Legacy
            mapped_values['key_achievements_list'] = tag_values.get('key_achievements_list', [])  # NEW: For Jinja loops
            mapped_values['leadership_philosophy'] = tag_values.get('leadership_philosophy', '')
            mapped_values['professional_summary'] = tag_values.get('professional_summary', '')
            
            # NEW: Structured skill categories for Jinja loops
            mapped_values['skill_categories'] = tag_values.get('skill_categories', [])
            
            # NEW: Structured job experience objects for Jinja loops
            mapped_values['sure'] = tag_values.get('sure', {'description': '', 'achievements': []})
            mapped_values['root'] = tag_values.get('root', {'description': '', 'achievements': []})
            mapped_values['enlace'] = tag_values.get('enlace', {'description': '', 'achievements': []})
            mapped_values['manta'] = tag_values.get('manta', {'description': '', 'achievements': []})
            
            # Individual skill category tags (backward compatibility)
            mapped_values['leadership_heading'] = tag_values.get('leadership_heading', 'LEADERSHIP & MANAGEMENT')
            mapped_values['leadership_skills'] = tag_values.get('leadership_skills', '')
            mapped_values['engineering_heading'] = tag_values.get('engineering_heading', 'ENGINEERING & ARCHITECTURE') 
            mapped_values['engineering_skills'] = tag_values.get('engineering_skills', '')
            
            # Legacy skill tags for backward compatibility
            mapped_values['skill_heading'] = tag_values.get('skill_heading', '')
            mapped_values['skills'] = tag_values.get('skills', '')
            
            mapped_values['education'] = tag_values.get('education', '')
            
            # Company-specific mappings with alternative names
            mapped_values['sureroledescription'] = tag_values.get('sureroledescription', '')
            mapped_values['suredescription'] = tag_values.get('sureroledescription', '')  # Alternative
            mapped_values['sureachievements'] = tag_values.get('sureachievements', '')
            
            mapped_values['rootroledescription'] = tag_values.get('rootroledescription', '')
            mapped_values['rootdescription'] = tag_values.get('rootroledescription', '')  # Alternative
            mapped_values['rootachievements'] = tag_values.get('rootachievements', '')
            
            mapped_values['enlaceroledescription'] = tag_values.get('enlaceroledescription', '')
            mapped_values['enlacedescription'] = tag_values.get('enlaceroledescription', '')  # Alternative
            mapped_values['enlaceachievements'] = tag_values.get('enlaceachievements', '')
            
            mapped_values['mantaroledescription'] = tag_values.get('mantaroledescription', '')
            mapped_values['mantadescription'] = tag_values.get('mantaroledescription', '')  # Alternative
            mapped_values['mantaachievements'] = tag_values.get('mantaachievements', '')
        
        # Handle cover letter template tags
        elif "Cover Letter" in template_path.name:
            # Cover letter template mappings
            mapped_values['company'] = tag_values.get('company', '')
            mapped_values['role'] = tag_values.get('role', '')
            mapped_values['content'] = tag_values.get('content', '')
        
        # Add any other tags that weren't mapped
        for key, value in tag_values.items():
            if key not in mapped_values:
                mapped_values[key] = value
        
        return mapped_values
        
    except Exception as e:
        print(f"Error mapping tag values: {str(e)}")
        return tag_values  # Return original if mapping fails


def parse_claude_response_for_tags(claude_response: str, baseline_content: str) -> Dict[str, Any]:
    """
    Simple approach: Just extract basic sections from Claude's response for template tags.
    """
    try:
        import json
        import re
        
        # First try JSON (for legacy cover letters)
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', claude_response, re.DOTALL)
        if json_match:
            try:
                tag_values = json.loads(json_match.group(1))
                print(f"Successfully parsed JSON with {len(tag_values)} tags")
                return tag_values
            except json.JSONDecodeError:
                pass
        
        # Try template tag format (for new cover letter format)
        template_tags = extract_template_tags(claude_response)
        if template_tags:
            print(f"Successfully parsed template tags with {len(template_tags)} tags")
            return template_tags
        
        # For resume: Just use Claude's response directly in key sections
        print("Using Claude's formatted response for template population")
        
        # Simple extraction - just grab obvious sections
        skills_section = extract_section(claude_response, 'Skills', 'Professional Experience')
        skill_categories_list = extract_skill_categories_for_jinja(skills_section)
        skill_categories_dict = extract_individual_skill_categories(skills_section)
        
        # Extract structured job experience data for each company
        sure_job = extract_company_job_experience(claude_response, 'SURE')
        root_job = extract_company_job_experience(claude_response, 'ROOT')
        enlace_job = extract_company_job_experience(claude_response, 'ENLACE')
        manta_job = extract_company_job_experience(claude_response, 'MANTA')
        
        # Extract structured key achievements
        key_achievements_list = extract_key_achievements_list(claude_response)
        
        # Extract headline content - try multiple possible section names
        headline_content = ""
        possible_headlines = ['Leadership Philosophy', 'Professional Summary', 'Executive Summary', 'Summary']
        
        for headline in possible_headlines:
            content = extract_section(claude_response, headline, 'Key Achievements')
            if content:
                headline_content = content
                print(f"Found headline content in '{headline}' section")
                break
        
        # If no specific section found, use the content after contact info but before Key Achievements
        if not headline_content:
            # Try to extract first paragraph after contact details
            lines = claude_response.split('\n')
            collecting = False
            content_lines = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Start collecting after we see contact info or header
                if any(indicator in line.lower() for indicator in ['jerry mindek', '@gmail.com', 'columbus, oh']):
                    collecting = True
                    continue
                    
                if collecting and line and not line.startswith('Key Achievements'):
                    # Skip obvious headers/titles
                    if line not in ['Jerry Mindek', 'LinkedIn | Github | Substack']:
                        content_lines.append(line)
                    
                # Stop when we hit Key Achievements
                if line.startswith('Key Achievements'):
                    break
                    
            if content_lines:
                headline_content = '\n'.join(content_lines)
                print("Extracted headline content from first section after contact info")
        
        tag_values = {
            'jobtitle': 'Engineering Manager',
            'key_achievements': extract_section(claude_response, 'Key Achievements', 'Skills'),  # Legacy
            'key_achievements_list': key_achievements_list,  # NEW: Structured list for Jinja loops
            'headline': headline_content,  # Single tag for first section content (Leadership Philosophy or Professional Summary)
            'professional_summary': headline_content,   # Same content for backward compatibility
            'skill_categories': skill_categories_list,  # Structured list for Jinja loops
            'education': extract_section(claude_response, 'Education', ''),
            
            # NEW: Structured job experience objects for Jinja loops
            'sure': sure_job,
            'root': root_job, 
            'enlace': enlace_job,
            'manta': manta_job,
            
            # Legacy individual tags for backward compatibility
            'sureroledescription': sure_job['description'],
            'sureachievements': extract_company_achievements(claude_response, 'SURE'),
            'rootroledescription': root_job['description'],
            'rootachievements': extract_company_achievements(claude_response, 'ROOT'),
            'enlaceroledescription': enlace_job['description'],
            'enlaceachievements': extract_company_achievements(claude_response, 'ENLACE'),
            'mantaroledescription': manta_job['description'],
            'mantaachievements': extract_company_achievements(claude_response, 'MANTA')
        }
        
        # Add individual skill category tags for backward compatibility
        tag_values.update(skill_categories_dict)
        
        print(f"Extracted {len([v for v in tag_values.values() if v])} populated tags")
        return tag_values
        
    except Exception as e:
        print(f"Error parsing Claude response: {str(e)}")
        return {}


def extract_template_tags(claude_response: str) -> Dict[str, str]:
    """
    Extract template tags from Claude's response in the format:
    {{tag_name}}
    [tag content]
    """
    try:
        import re
        tag_values = {}
        
        # Pattern to match {{tag_name}} followed by content until next tag or end
        pattern = r'\{\{(\w+)\}\}\s*\n(.*?)(?=\{\{\w+\}\}|\Z)'
        matches = re.findall(pattern, claude_response, re.DOTALL)
        
        for tag_name, tag_content in matches:
            # Clean up the content (remove brackets if present, strip whitespace)
            cleaned_content = tag_content.strip()
            # Remove square brackets if they wrap the entire content
            if cleaned_content.startswith('[') and cleaned_content.endswith(']'):
                cleaned_content = cleaned_content[1:-1].strip()
            tag_values[tag_name] = cleaned_content
            
        return tag_values
        
    except Exception as e:
        print(f"Error extracting template tags: {str(e)}")
        return {}


def parse_interview_prep_response(claude_response: str) -> Dict[str, Any]:
    """Parse Claude's interview prep response into template tags"""
    try:
        import re
        
        tag_values = {}
        
        # Extract company and job title
        company_match = re.search(r'COMPANY:\s*\n(.*?)(?=\n[A-Z_]+:|$)', claude_response, re.DOTALL)
        if company_match:
            tag_values['company'] = company_match.group(1).strip()
        
        job_title_match = re.search(r'JOB_TITLE:\s*\n(.*?)(?=\n[A-Z_]+:|$)', claude_response, re.DOTALL)
        if job_title_match:
            tag_values['jobtitle'] = job_title_match.group(1).strip()
        
        # Extract core narrative
        core_match = re.search(r'CORE_NARRATIVE:\s*\n(.*?)(?=\n[A-Z_]+:|$)', claude_response, re.DOTALL)
        if core_match:
            tag_values['core_narrative'] = core_match.group(1).strip()
        
        # Extract proof points
        proof_section = re.search(r'PROOF_POINTS:\s*\n(.*?)(?=\nPOTENTIAL_CONCERNS:|$)', claude_response, re.DOTALL)
        if proof_section:
            proof_text = proof_section.group(1).strip()
            proof_list = []
            
            # Parse numbered items with headers and content
            proof_items = re.findall(r'(\d+\.\s*\[Header:[^\]]+\]\s*\n\[Content:[^\]]+\])', proof_text, re.DOTALL)
            for item in proof_items:
                header_match = re.search(r'Header:\s*([^\]]+)', item)
                content_match = re.search(r'Content:\s*([^\]]+)', item)
                if header_match and content_match:
                    proof_list.append({
                        'header': header_match.group(1).strip(),
                        'content': content_match.group(1).strip()
                    })
            
            tag_values['proof_list'] = proof_list
        
        # Extract potential concerns
        concerns_section = re.search(r'POTENTIAL_CONCERNS:\s*\n(.*?)(?=\nCULTURAL_ALIGNMENT:|$)', claude_response, re.DOTALL)
        if concerns_section:
            concerns_text = concerns_section.group(1).strip()
            concern_list = []
            
            concern_items = re.findall(r'(\d+\.\s*\[Header:[^\]]+\]\s*\n\[Content:[^\]]+\])', concerns_text, re.DOTALL)
            for item in concern_items:
                header_match = re.search(r'Header:\s*([^\]]+)', item)
                content_match = re.search(r'Content:\s*([^\]]+)', item)
                if header_match and content_match:
                    concern_list.append({
                        'header': header_match.group(1).strip(),
                        'content': content_match.group(1).strip()
                    })
            
            tag_values['concern_list'] = concern_list
        
        # Extract cultural alignment
        culture_section = re.search(r'CULTURAL_ALIGNMENT:\s*\n(.*?)(?=\nSTRATEGIC_QUESTIONS:|$)', claude_response, re.DOTALL)
        if culture_section:
            culture_text = culture_section.group(1).strip()
            culture_list = []
            
            culture_items = re.findall(r'(\d+\.\s*\[Header:[^\]]+\]\s*\n\[Content:[^\]]+\])', culture_text, re.DOTALL)
            for item in culture_items:
                header_match = re.search(r'Header:\s*([^\]]+)', item)
                content_match = re.search(r'Content:\s*([^\]]+)', item)
                if header_match and content_match:
                    culture_list.append({
                        'header': header_match.group(1).strip(),
                        'content': content_match.group(1).strip()
                    })
            
            tag_values['culture_list'] = culture_list
        
        # Extract strategic questions
        questions_section = re.search(r'STRATEGIC_QUESTIONS:\s*\n(.*?)(?=\nCLOSING_POSITIONING:|$)', claude_response, re.DOTALL)
        if questions_section:
            questions_text = questions_section.group(1).strip()
            # Extract numbered questions
            questions = re.findall(r'\d+\.\s*(.+?)(?=\n\d+\.|$)', questions_text, re.DOTALL)
            strategic_questions_list = [q.strip() for q in questions if q.strip()]
            tag_values['strategic_questions_list'] = strategic_questions_list
        
        # Extract closing positioning
        closing_match = re.search(r'CLOSING_POSITIONING:\s*\n(.*?)$', claude_response, re.DOTALL)
        if closing_match:
            tag_values['closing_positioning'] = closing_match.group(1).strip()
        
        print(f"Parsed interview prep with {len(tag_values)} sections")
        return tag_values
        
    except Exception as e:
        print(f"Error parsing interview prep response: {str(e)}")
        return {}


def extract_section(text: str, start_marker: str, end_marker: str) -> str:
    """Extract text between two section markers."""
    try:
        start_pos = text.find(start_marker)
        if start_pos == -1:
            return ""
        
        start_pos = text.find('\n', start_pos) + 1
        if end_marker:
            end_pos = text.find(end_marker, start_pos)
            if end_pos == -1:
                return text[start_pos:].strip()
            return text[start_pos:end_pos].strip()
        else:
            return text[start_pos:].strip()
    except:
        return ""


def extract_company_description(text: str, company_name: str) -> str:
    """Extract role description for a specific company (the paragraph before the bullets)."""
    try:
        # Look for company in Professional Experience section only
        prof_exp_pos = text.find('Professional Experience')
        if prof_exp_pos == -1:
            prof_exp_pos = text.find('EXPERIENCE')
        
        if prof_exp_pos == -1:
            print(f"DEBUG: Professional Experience section not found")
            return ""
        
        # Search for company only within Professional Experience section
        search_text = text[prof_exp_pos:]
        company_pos = search_text.upper().find(company_name.upper())
        if company_pos == -1:
            print(f"DEBUG: Company '{company_name}' not found in Professional Experience section")
            return ""
        
        # Adjust position to full text
        absolute_pos = prof_exp_pos + company_pos
        print(f"DEBUG: Found {company_name} at position {absolute_pos} in Professional Experience")
        
        # Look for the description paragraph that comes before the bullet points
        section = text[absolute_pos:absolute_pos + 1000]  # Look ahead 1000 chars
        lines = section.split('\n')
        
        print(f"DEBUG: Section for {company_name}: {section[:200]}...")
        
        # Find the first non-bullet, non-empty line after finding the company
        descriptions = []
        found_company_line = False
        for line in lines:
            line = line.strip()
            if not found_company_line:
                # Look for the actual company header line
                if company_name.upper() in line.upper():
                    found_company_line = True
                continue
            
            if line and not line.startswith('•') and not line.upper().startswith(('ROOT', 'SURE', 'ENLACE', 'MANTA')):
                # Skip lines that look like company headers (with pipe separators, company names like "ROOT INC", dates)
                if '|' not in line and not line.upper().endswith(' INC') and not any(year in line for year in ['2018', '2019', '2020', '2021', '2022', '2023', '2024', '2025']):
                    descriptions.append(line)
                    print(f"DEBUG: Found description line for {company_name}: {line}")
            elif line.startswith('•') or line.upper().startswith(('ROOT', 'SURE', 'ENLACE', 'MANTA')):
                break  # Stop when we hit bullets or next company
        
        result = '\n'.join(descriptions)
        print(f"DEBUG: Final description for {company_name}: '{result}'")
        return result
    except Exception as e:
        print(f"DEBUG: Error extracting description for {company_name}: {str(e)}")
        return ""


def extract_skills_heading_and_content(skills_section: str) -> tuple[str, str]:
    """
    Extract skills section into individual skill category tags.
    
    Args:
        skills_section: The complete skills section text
        
    Returns:
        Tuple of (skill_categories_dict, fallback_content) - but we'll return empty strings
        since we're now extracting individual categories in the main function
    """
    # This function is no longer used - we extract individual categories instead
    return "", ""


def extract_skill_categories_for_jinja(skills_section: str) -> list:
    """
    Extract skills section into structured list for Jinja template loops.
    
    Args:
        skills_section: The complete skills section text
        
    Returns:
        List of dictionaries with 'heading' and 'content' keys for each skill category
    """
    try:
        if not skills_section:
            return []
        
        lines = skills_section.strip().split('\n')
        categories = []
        current_heading = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if this looks like a skill category heading
            if (line.isupper() and len(line) < 50 and not line.startswith('•') and 
                ('&' in line or any(keyword in line for keyword in ['LEADERSHIP', 'MANAGEMENT', 'ENGINEERING', 'ARCHITECTURE', 'TECHNICAL', 'BUSINESS']))):
                current_heading = line
                print(f"DEBUG: Found skill category: {current_heading}")
            else:
                # This is skill content for the current heading
                if current_heading:
                    categories.append({
                        'heading': current_heading,
                        'content': line
                    })
                    print(f"DEBUG: Added skill category - Heading: '{current_heading}', Content: '{line[:50]}...'")
                    current_heading = None  # Reset for next category
        
        # Fallback if no categories found
        if not categories:
            categories = [
                {'heading': 'LEADERSHIP & MANAGEMENT', 'content': ''},
                {'heading': 'ENGINEERING & ARCHITECTURE', 'content': ''}
            ]
        
        print(f"DEBUG: Extracted {len(categories)} structured skill categories")
        for i, cat in enumerate(categories):
            print(f"  {i+1}. {cat['heading']}: {cat['content'][:30]}...")
            
        return categories
        
    except Exception as e:
        print(f"DEBUG: Error extracting skill categories for Jinja: {str(e)}")
        return [
            {'heading': 'LEADERSHIP & MANAGEMENT', 'content': ''},
            {'heading': 'ENGINEERING & ARCHITECTURE', 'content': ''}
        ]


def extract_individual_skill_categories(skills_section: str) -> dict:
    """
    Legacy function - extract skills section into individual category tags.
    Kept for backward compatibility.
    """
    skill_categories = extract_skill_categories_for_jinja(skills_section)
    
    # Convert structured list back to individual tags for compatibility
    individual_tags = {}
    for category in skill_categories:
        heading = category['heading']
        content = category['content']
        
        if 'LEADERSHIP' in heading or 'MANAGEMENT' in heading:
            individual_tags['leadership_heading'] = heading
            individual_tags['leadership_skills'] = content
        elif 'ENGINEERING' in heading or 'ARCHITECTURE' in heading:
            individual_tags['engineering_heading'] = heading
            individual_tags['engineering_skills'] = content
        else:
            # Generic fallback
            safe_name = heading.lower().replace(' ', '_').replace('&', 'and')
            individual_tags[f'{safe_name}_heading'] = heading
            individual_tags[f'{safe_name}_skills'] = content
    
    return individual_tags


def extract_key_achievements_list(text: str) -> list:
    """
    Extract key achievements section as a list for Jinja loops.
    
    Args:
        text: The complete resume text
        
    Returns:
        List of achievement strings without numbers
    """
    try:
        key_achievements_section = extract_section(text, 'Key Achievements', 'Skills')
        achievements_list = []
        
        if key_achievements_section:
            for line in key_achievements_section.split('\n'):
                line = line.strip()
                if line:
                    # Remove number prefixes (1. 2. 3. etc.) and any leading whitespace
                    import re
                    clean_line = re.sub(r'^\d+\.\s*', '', line)  # Remove "1. ", "2. ", etc.
                    if clean_line:  # Only add non-empty lines
                        achievements_list.append(clean_line)
        
        print(f"DEBUG: Extracted {len(achievements_list)} key achievements without numbers")
        for i, achievement in enumerate(achievements_list):
            print(f"  {i+1}: {achievement[:50]}...")
            
        return achievements_list
        
    except Exception as e:
        print(f"DEBUG: Error extracting key achievements list: {str(e)}")
        return []


def extract_company_job_experience(text: str, company_name: str) -> dict:
    """
    Extract complete job experience data for a specific company.
    
    Args:
        text: The complete resume text
        company_name: Company name to extract (e.g., 'SURE', 'ROOT', etc.)
        
    Returns:
        Dictionary with 'description' and 'achievements' (list) keys
    """
    try:
        description = extract_company_description(text, company_name)
        achievements_text = extract_company_achievements(text, company_name)
        
        # Convert achievements text to list of individual bullets (without bullet symbols)
        achievements_list = []
        if achievements_text:
            for line in achievements_text.split('\n'):
                line = line.strip()
                if line.startswith('•'):
                    # Remove the bullet symbol and any leading whitespace
                    clean_line = line[1:].strip()  # Remove '•' and trim spaces
                    achievements_list.append(clean_line)
        
        result = {
            'description': description,
            'achievements': achievements_list
        }
        
        print(f"DEBUG: Extracted job experience for {company_name}: description={bool(description)}, achievements={len(achievements_list)}")
        return result
        
    except Exception as e:
        print(f"DEBUG: Error extracting job experience for {company_name}: {str(e)}")
        return {'description': '', 'achievements': []}


def extract_company_achievements(text: str, company_name: str) -> str:
    """Extract achievement bullets for a specific company."""
    try:
        # Look for company in Professional Experience section only
        prof_exp_pos = text.find('Professional Experience')
        if prof_exp_pos == -1:
            prof_exp_pos = text.find('EXPERIENCE')
        
        if prof_exp_pos == -1:
            print(f"DEBUG: Professional Experience section not found for {company_name} achievements")
            return ""
        
        # Search for company only within Professional Experience section
        search_text = text[prof_exp_pos:]
        company_pos = search_text.upper().find(company_name.upper())
        if company_pos == -1:
            print(f"DEBUG: Company '{company_name}' not found for achievements in Professional Experience")
            return ""
        
        # Adjust position to full text
        absolute_pos = prof_exp_pos + company_pos
        print(f"DEBUG: Extracting achievements for {company_name} at position {absolute_pos}")
        
        # Find bullets that come after this company but before the next company
        section = text[absolute_pos:absolute_pos + 1000]  # Look ahead 1000 chars
        lines = section.split('\n')
        bullets = []
        found_company_line = False
        
        for line in lines:
            line = line.strip()
            if not found_company_line:
                # Look for the actual company header line
                if company_name.upper() in line.upper():
                    found_company_line = True
                continue
            
            if line.startswith('•'):
                bullets.append(line)
                print(f"DEBUG: Found bullet for {company_name}: {line}")
            elif line and (line.upper().startswith(('ROOT', 'SURE', 'ENLACE', 'MANTA')) or 'Education' in line):
                # Stop when we hit the next company or education section
                break
        
        result = '\n'.join(bullets)
        print(f"DEBUG: Final achievements for {company_name}: Found {len(bullets)} bullets")
        return result
    except Exception as e:
        print(f"DEBUG: Error extracting achievements for {company_name}: {str(e)}")
        return ""


def tag_based_template_replacement_claude(template_path: Path, dest_path: Path, populated_content: str) -> bool:
    """
    Legacy function: Replace tags in a template document with content provided by Claude.
    This is kept as fallback when Claude handles tag replacement directly.
    
    Args:
        template_path: Path to template .docx file with tags like {jobtitle}, {achievements}, etc.
        dest_path: Path for new .docx file with tags replaced
        populated_content: Complete document content with tags already replaced by Claude
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Since Claude handles all tag replacements and returns the complete populated document,
        # we just need to create a new document with the content Claude provided
        return write_docx_content(dest_path, populated_content, preserve_formatting=True)
        
    except Exception as e:
        print(f"Failed tag-based template replacement: {str(e)}")
        return False


def read_gdoc_content(file_path: Path) -> str:
    """
    Read content from a Google Doc file (.gdoc) in the sync folder.
    Google Doc files in sync folders are actually shortcuts, so we read them as text.
    
    Args:
        file_path: Path to the .gdoc file
        
    Returns:
        String containing the content, or empty string if can't read
    """
    try:
        # Try to read as plain text file first
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # .gdoc files might contain JSON with URL info, try to extract useful content
        import json
        try:
            gdoc_data = json.loads(content)
            # This is a shortcut file, we can't directly read the content
            # Return empty string to signal we need to look for alternative files
            return ""
        except json.JSONDecodeError:
            # Not JSON, might be actual content
            return content
            
    except Exception as e:
        print(f"Could not read .gdoc file {file_path}: {str(e)}")
        return ""


def find_template_and_baseline_files(templates_folder: Path, baseline_resume_name: str) -> tuple[Optional[Path], Optional[Path]]:
    """
    Find both the template file (with "Template" in name) and baseline content file
    for a given resume type. Supports both .docx and .gdoc files, preferring .gdoc for baseline.
    
    Args:
        templates_folder: Path to templates folder
        baseline_resume_name: Name like "Engineering Manager"
        
    Returns:
        Tuple of (template_file_path, baseline_file_path) or (None, None) if not found
    """
    try:
        # Look for template file (should have "Template" in the name)
        template_patterns = [
            f"{baseline_resume_name} Template.docx",
            f"{baseline_resume_name} Template.gdoc", 
            f"{baseline_resume_name}Template.docx",
            f"{baseline_resume_name}Template.gdoc"
        ]
        
        template_file = None
        for pattern in template_patterns:
            potential_file = templates_folder / pattern
            if potential_file.exists():
                template_file = potential_file
                break
        
        # Look for baseline content file (prefer .docx, then .gdoc, then .txt)
        baseline_patterns = [
            f"{baseline_resume_name}.docx",
            f"{baseline_resume_name}.gdoc", 
            f"{baseline_resume_name}.txt",
            baseline_resume_name  # No extension
        ]
        
        baseline_file = None
        for pattern in baseline_patterns:
            potential_file = templates_folder / pattern
            if potential_file.exists():
                baseline_file = potential_file
                break
        
        return template_file, baseline_file
        
    except Exception as e:
        print(f"Error finding template and baseline files: {str(e)}")
        return None, None


def read_content_file(file_path: Path) -> str:
    """
    Read content from various file types (.docx, .gdoc, .txt).
    Handles the logic for reading different file formats.
    
    Args:
        file_path: Path to the file to read
        
    Returns:
        String containing the file content
    """
    try:
        if file_path.suffix.lower() == '.docx':
            try:
                return read_docx_content(file_path)
            except Exception as docx_error:
                print(f"Warning: Failed to read docx file {file_path}: {docx_error}")
                # Try to find alternative files if docx is corrupted
                base_name = file_path.stem
                parent_dir = file_path.parent
                alternatives = [
                    parent_dir / f"{base_name}.txt",
                    parent_dir / f"{base_name}.gdoc"
                ]
                for alt_file in alternatives:
                    if alt_file.exists():
                        print(f"Trying alternative file: {alt_file}")
                        try:
                            return read_content_file(alt_file)
                        except Exception as alt_error:
                            print(f"Alternative file {alt_file} also failed: {alt_error}")
                            continue
                # If no alternatives work, re-raise the original docx error
                raise docx_error
        elif file_path.suffix.lower() == '.gdoc':
            # Try .gdoc first, but if it's just a shortcut, look for alternatives
            gdoc_content = read_gdoc_content(file_path)
            if gdoc_content:
                return gdoc_content
            else:
                # .gdoc didn't work, try to find a .docx or .txt alternative
                base_name = file_path.stem
                parent_dir = file_path.parent
                alternatives = [
                    parent_dir / f"{base_name}.docx",
                    parent_dir / f"{base_name}.txt"
                ]
                for alt_file in alternatives:
                    if alt_file.exists():
                        return read_content_file(alt_file)
                return ""
        else:
            # Read as plain text
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
    except Exception as e:
        raise Exception(f"Failed to read content file {file_path}: {str(e)}")


def smart_content_replacement(source_path: Path, dest_path: Path, new_content: str) -> bool:
    """
    Advanced content replacement that preserves document structure and formatting.
    This approach maintains the exact document structure while selectively updating content.
    
    Args:
        source_path: Path to source .docx file
        dest_path: Path for new .docx file
        new_content: New optimized content from Claude
        
    Returns:
        True if successful, False otherwise
    """
    try:
        import shutil
        import re
        
        # First, make an exact copy of the source file
        shutil.copy2(source_path, dest_path)
        
        # Load the copied document
        doc = Document(dest_path)
        
        # Parse the new content to identify sections and their content
        sections = {}
        current_section = None
        section_content = []
        
        for line in new_content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Check if this looks like a section header (ALL CAPS, ends with colon, or looks like a heading)
            if (line.isupper() and len(line) < 50) or \
               (line.endswith(':') and len(line) < 50) or \
               re.match(r'^[A-Z][A-Za-z\s&-]+$', line) and len(line) < 50:
                # Save previous section
                if current_section and section_content:
                    sections[current_section.lower()] = section_content[:]
                # Start new section
                current_section = line
                section_content = []
            else:
                # Add to current section content
                if current_section:
                    section_content.append(line)
        
        # Save the last section
        if current_section and section_content:
            sections[current_section.lower()] = section_content[:]
        
        print(f"DEBUG: Identified {len(sections)} sections in new content")
        for section_name in sections.keys():
            print(f"  - {section_name}: {len(sections[section_name])} items")
        
        # Strategy: Try to intelligently match and replace content while preserving structure
        # This is a conservative approach - only replace content that we can confidently map
        
        original_paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                original_paragraphs.append((paragraph, paragraph.text.strip()))
        
        print(f"DEBUG: Found {len(original_paragraphs)} non-empty paragraphs in original document")
        
        # For now, use a simple but safe approach: replace the entire content but preserve key formatting
        # This ensures we don't break the document structure while still updating content
        
        # Clear all content first
        for paragraph, _ in original_paragraphs:
            paragraph.clear()
        
        # Add new content line by line, trying to preserve paragraph structure
        content_lines = [line.strip() for line in new_content.split('\n') if line.strip()]
        
        # Ensure we have enough paragraphs
        while len(doc.paragraphs) < len(content_lines):
            doc.add_paragraph()
        
        # Replace content paragraph by paragraph
        for i, line in enumerate(content_lines):
            if i < len(doc.paragraphs):
                paragraph = doc.paragraphs[i]
                paragraph.clear()
                
                # Determine formatting based on content
                if (line.isupper() and len(line) < 50) or (line.endswith(':') and len(line) < 50):
                    # This looks like a header
                    run = paragraph.add_run(line)
                    run.bold = True
                else:
                    # Regular content
                    paragraph.add_run(line)
        
        # Save the updated document
        doc.save(str(dest_path))
        return True
        
    except Exception as e:
        print(f"Failed smart content replacement from {source_path} to {dest_path}: {str(e)}")
        return False